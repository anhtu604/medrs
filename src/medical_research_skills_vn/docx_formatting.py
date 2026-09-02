"""Apply source-traceable Word mechanics to a new DOCX artifact."""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .structure_profiles import ROOT, _sha256, verify_profile_source


def _field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def _add_introduction_section(document, heading: str):
    target = heading.casefold().strip()
    index = next((i for i, p in enumerate(document.paragraphs) if p.text.casefold().strip() == target), None)
    if index is None or index == 0:
        return False
    preceding = document.paragraphs[index - 1]
    ppr = preceding._p.get_or_add_pPr()
    old = ppr.find(qn("w:sectPr"))
    if old is None:
        ppr.append(deepcopy(document.element.body.sectPr))
    body_sectpr = document.element.body.sectPr
    pg_num = body_sectpr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        body_sectpr.append(pg_num)
    pg_num.set(qn("w:start"), "1")
    return True


def format_docx(source_path: Path, output_path: Path, profile: dict, *, author_approved: bool = False) -> dict:
    source = Path(source_path).resolve()
    output = Path(output_path).resolve()
    if source == output:
        return {"status": "SOURCE_OVERWRITE_FORBIDDEN"}
    verification = verify_profile_source(ROOT, profile)
    if verification["status"] != "VERIFIED":
        return verification
    if not author_approved:
        return {"status": "AUTHOR_APPROVAL_REQUIRED"}

    document = Document(source)
    body = profile["body"]
    normal = document.styles["Normal"]
    normal.font.name = body["font"]
    normal.font.size = Pt(body["size_pt"])
    normal.paragraph_format.line_spacing = body["line_spacing"]
    for style in document.styles:
        if getattr(style, "font", None) is not None:
            style.font.name = body["font"]
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = body["font"]
            run.font.size = Pt(body["size_pt"])

    page = profile["page"]
    for section in document.sections:
        section.top_margin = Cm(page["top_cm"])
        section.bottom_margin = Cm(page["bottom_cm"])
        section.left_margin = Cm(page["left_cm"])
        section.right_margin = Cm(page["right_cm"])

    toc_heading = next((p for p in document.paragraphs if p.text.casefold().strip() == "mục lục"), None)
    if toc_heading is not None:
        toc_paragraph = OxmlElement("w:p")
        toc = OxmlElement("w:fldSimple")
        toc.set(qn("w:instr"), 'TOC \\o "1-4" \\h \\z \\u')
        toc_paragraph.append(toc)
        toc_heading._p.addnext(toc_paragraph)

    if not _add_introduction_section(document, profile["numbering"]["restart_section_heading"]):
        return {"status": "AUTHOR_INPUT_REQUIRED", "reason": "INTRODUCTION_HEADING_NOT_FOUND"}
    sections = document.sections
    body_section = sections[-1]
    body_section.header.is_linked_to_previous = False
    header_paragraph = body_section.header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(header_paragraph, "PAGE")

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    return {
        "status": "FORMATTED_WITH_UNPERFORMED_CHECKS",
        "backend": "ooxml-only",
        "source_hash": _sha256(source),
        "output_hash": _sha256(output),
        "output": str(output),
        "unperformed_checks": ["field-update", "pagination", "visual-layout"],
    }
