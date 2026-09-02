"""Source-traceable semantic document restructuring."""

from __future__ import annotations

import hashlib
import re
import unicodedata
from copy import deepcopy
from datetime import date
from pathlib import Path

import yaml
from docx import Document
from docx.text.paragraph import Paragraph

from .freshness import verification_status


ROOT = Path(__file__).parents[2]


def load_structure_profile(path: Path) -> dict:
    return yaml.safe_load(Path(path).read_text(encoding="utf-8")) or {}


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with Path(path).open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def verify_profile_source(root: Path, profile: dict) -> dict:
    if profile.get("verification_status") != "CURRENT" or verification_status(profile, date.today()) != "CURRENT":
        return {"status": "OFFICIAL_RULE_REQUIRED", "reason": "PROFILE_STALE_OR_UNVERIFIED"}
    source = Path(root) / profile["source_path"]
    if not source.is_file():
        if profile.get("source_url") and profile.get("source_sha256"):
            return {
                "status": "VERIFIED",
                "verification_basis": "REGISTERED_PROVENANCE_SOURCE_NOT_BUNDLED",
                "source_available": False,
                "sha256": profile["source_sha256"].upper(),
                "source_url": profile["source_url"],
            }
        return {"status": "OFFICIAL_RULE_REQUIRED", "reason": "SOURCE_MISSING"}
    actual_hash = _sha256(source)
    if actual_hash != profile.get("source_sha256", "").upper():
        return {"status": "OFFICIAL_RULE_REQUIRED", "reason": "SOURCE_HASH_MISMATCH", "sha256": actual_hash}
    document = Document(source)
    for rule in profile.get("semantic_rules", []):
        match = re.fullmatch(r"P(\d+)", str(rule.get("source_locator", "")))
        if not match or int(match.group(1)) >= len(document.paragraphs) or not rule.get("summary"):
            return {"status": "OFFICIAL_RULE_REQUIRED", "reason": "SOURCE_LOCATOR_INVALID", "rule": rule.get("id")}
    return {
        "status": "VERIFIED",
        "verification_basis": "LOCAL_SOURCE_HASH_AND_LOCATORS",
        "source_available": True,
        "sha256": actual_hash,
        "source": str(source),
    }


def _normalize(text: str) -> str:
    text = unicodedata.normalize("NFD", text.casefold())
    text = "".join(char for char in text if unicodedata.category(char) != "Mn")
    return " ".join(re.sub(r"[^a-z0-9]+", " ", text).split())


def _alias_map(profile: dict) -> dict[str, str]:
    aliases = {}
    for section in profile.get("sections", []):
        candidates = [section.get("label", ""), *section.get("aliases", [])]
        for candidate in candidates:
            aliases[_normalize(candidate)] = section["id"]
    return aliases


def _recognize(text: str, aliases: dict[str, str]) -> str | None:
    normalized = _normalize(text)
    if normalized in aliases:
        return aliases[normalized]
    # Accept numbered headings after punctuation normalization, but not arbitrary body prose.
    for alias, section_id in sorted(aliases.items(), key=lambda item: len(item[0]), reverse=True):
        if len(alias) >= 8 and normalized.startswith(alias):
            return section_id
    return None


def inspect_docx_structure(path: Path, profile: dict) -> dict:
    document = Document(path)
    aliases = _alias_map(profile)
    recognized = []
    headings = []
    for paragraph in document.paragraphs:
        style_name = paragraph.style.name.casefold() if paragraph.style else ""
        if not style_name.startswith("heading") and not style_name.startswith("tiêu đề"):
            continue
        section_id = _recognize(paragraph.text, aliases)
        if section_id:
            recognized.append(section_id)
            headings.append({"section_id": section_id, "text": paragraph.text, "style": paragraph.style.name})
    expected = [item["id"] for item in profile.get("sections", []) if item["id"] in recognized]
    return {
        "status": "PASS" if recognized == expected else "FAIL",
        "recognized": recognized,
        "expected_present_order": expected,
        "headings": headings,
    }


def plan_restructure(current_sections: list[str], profile: dict, author_approved: bool = False) -> dict:
    ordered = [item["id"] for item in profile.get("sections", [])]
    present_order = [item for item in ordered if item in current_sections]
    required = {
        item["id"] for item in profile.get("sections", []) if item.get("required_for_restructure", False)
    }
    missing = sorted(required - set(current_sections), key=lambda item: ordered.index(item))
    moves = [
        {"section": section, "from": current_sections.index(section), "to": present_order.index(section)}
        for section in present_order
        if current_sections.index(section) != present_order.index(section)
    ]
    status = "AUTHOR_INPUT_REQUIRED" if missing else "READY"
    if moves and not author_approved and not missing:
        status = "AUTHOR_APPROVAL_REQUIRED"
    return {"status": status, "target_order": present_order, "moves": moves, "missing_required": missing}


def _recognized_groups(document, profile: dict):
    aliases = _alias_map(profile)
    body = document.element.body
    children = list(body.iterchildren())
    starts = []
    for index, child in enumerate(children):
        if not child.tag.endswith("}p"):
            continue
        paragraph = Paragraph(child, document)
        style_name = paragraph.style.name.casefold() if paragraph.style else ""
        if not style_name.startswith("heading") and not style_name.startswith("tiêu đề"):
            continue
        section_id = _recognize(paragraph.text, aliases)
        if section_id:
            starts.append((index, section_id))
    groups = {}
    for position, (start, section_id) in enumerate(starts):
        end = starts[position + 1][0] if position + 1 < len(starts) else len(children)
        if end and children[end - 1].tag.endswith("}sectPr"):
            end -= 1
        groups[section_id] = children[start:end]
    return groups


def apply_docx_structure(source_path: Path, output_path: Path, profile: dict, author_approved: bool = False) -> dict:
    source_path = Path(source_path).resolve()
    output_path = Path(output_path).resolve()
    if source_path == output_path:
        return {"status": "SOURCE_OVERWRITE_FORBIDDEN"}

    source_verification = verify_profile_source(ROOT, profile)
    if source_verification["status"] != "VERIFIED":
        return source_verification

    before = inspect_docx_structure(source_path, profile)
    duplicate_sections = sorted(
        {section for section in before["recognized"] if before["recognized"].count(section) > 1}
    )
    if duplicate_sections:
        return {"status": "AUTHOR_INPUT_REQUIRED", "duplicate_sections": duplicate_sections}
    plan = plan_restructure(before["recognized"], profile, author_approved=author_approved)
    if plan["missing_required"]:
        return {"status": "AUTHOR_INPUT_REQUIRED", "missing_required": plan["missing_required"], "plan": plan}
    if plan["moves"] and not author_approved:
        return {"status": "AUTHOR_APPROVAL_REQUIRED", "plan": plan}

    document = Document(source_path)
    groups = _recognized_groups(document, profile)
    body = document.element.body
    for elements in groups.values():
        for element in elements:
            if element.getparent() is body:
                body.remove(element)
    section_properties = body.sectPr
    insertion_index = body.index(section_properties) if section_properties is not None else len(body)
    for section_id in plan["target_order"]:
        for element in groups[section_id]:
            body.insert(insertion_index, deepcopy(element))
            insertion_index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    validation = inspect_docx_structure(output_path, profile)
    return {
        "status": "STRUCTURE_APPLIED" if validation["status"] == "PASS" else "VALIDATION_FAILED",
        "source": str(source_path),
        "output": str(output_path),
        "source_hash": _sha256(source_path),
        "output_hash": _sha256(output_path),
        "changes": plan["moves"],
        "validation": validation,
        "missing_required": [],
    }
