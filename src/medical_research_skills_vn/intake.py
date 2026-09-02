"""Extract an in-progress research project into a draft Passport."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

from docx import Document

from .passport import new_passport, unresolved_field


@dataclass(frozen=True)
class ExtractedBlock:
    text: str
    heading: str | None
    location: dict
    kind: str = "paragraph"


@dataclass(frozen=True)
class ExtractedProject:
    path: Path
    blocks: tuple[ExtractedBlock, ...]
    tables: tuple[tuple[tuple[str, ...], ...], ...]


def _extract_docx(path: Path) -> ExtractedProject:
    document = Document(path)
    blocks: list[ExtractedBlock] = []
    current_heading: str | None = None
    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        if style.startswith("Heading") or style == "Title":
            current_heading = text
            kind = "title" if style == "Title" else "heading"
        else:
            kind = "paragraph"
        blocks.append(ExtractedBlock(text, current_heading, {"paragraph": index}, kind))
    tables = []
    for table in document.tables:
        tables.append(tuple(tuple(cell.text.strip() for cell in row.cells) for row in table.rows))
    return ExtractedProject(path, tuple(blocks), tuple(tables))


def _extract_markdown(path: Path) -> ExtractedProject:
    blocks: list[ExtractedBlock] = []
    current_heading: str | None = None
    for line_number, raw in enumerate(path.read_text(encoding="utf-8").splitlines(), start=1):
        text = raw.strip()
        if not text:
            continue
        if text.startswith("#"):
            current_heading = text.lstrip("#").strip()
            kind = "title" if text.startswith("# ") else "heading"
            value = current_heading
        else:
            kind = "paragraph"
            value = text
        blocks.append(ExtractedBlock(value, current_heading, {"line": line_number}, kind))
    return ExtractedProject(path, tuple(blocks), ())


def extract_project(path):
    path = Path(path)
    suffix = path.suffix.casefold()
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix in {".md", ".markdown", ".txt"}:
        return _extract_markdown(path)
    raise ValueError(f"Unsupported intake format: {suffix}")


def _provenance(project: ExtractedProject, block: ExtractedBlock) -> dict:
    return {"artifact": str(project.path), **block.location, "excerpt": block.text[:240]}


def _inferred(project: ExtractedProject, candidates: list[ExtractedBlock]) -> dict:
    unique = []
    for candidate in candidates:
        if candidate.text not in [item.text for item in unique]:
            unique.append(candidate)
    value = [item.text for item in unique] if len(unique) > 1 else unique[0].text
    return {
        "value": value,
        "status": "DRAFT_INFERRED",
        "confidence": 0.75 if len(unique) == 1 else 0.45,
        "provenance": [_provenance(project, item) for item in unique],
        "confirmed_by": None,
        "marker": "AUTHOR_APPROVAL_REQUIRED",
    }


def draft_passport_from_project(project, project_id):
    passport = new_passport(project_id, "unknown-existing-project", "unresolved-locale")
    title_candidates = [block for block in project.blocks if block.kind == "title"][:1]
    objective_candidates = [
        block
        for block in project.blocks
        if block.kind == "paragraph"
        and ("mục tiêu" in block.text.casefold() or "mục tiêu" in (block.heading or "").casefold())
    ]
    population_candidates = [
        block for block in project.blocks if "đối tượng nghiên cứu" in block.text.casefold()
    ]
    design_pattern = re.compile(r"\b(cắt ngang|thuần tập|bệnh.chứng|ngẫu nhiên|cohort|case.control|cross.sectional|rct)\b", re.I)
    design_candidates = [block for block in project.blocks if design_pattern.search(block.text)]

    passport["title"] = _inferred(project, title_candidates) if title_candidates else unresolved_field()
    if objective_candidates:
        passport["objectives"] = _inferred(project, objective_candidates)
    if population_candidates:
        passport["population"] = _inferred(project, population_candidates)
    if design_candidates:
        passport["study_design"] = _inferred(project, design_candidates)
    passport["ethics_approval"] = unresolved_field("AUTHOR_APPROVAL_REQUIRED")
    passport["intake_questions"] = []
    if len({item.text for item in objective_candidates}) > 1:
        passport["intake_questions"].append(
            {
                "field": "objectives",
                "question_vi": "Bản thảo có nhiều cách nêu mục tiêu. Mục tiêu nào là bản đã được tác giả duyệt?",
                "question_en": "The draft contains conflicting objective statements. Which version has author approval?",
            }
        )
    for field, label in (("study_design", "thiết kế nghiên cứu"), ("ethics_approval", "phê duyệt đạo đức")):
        if passport[field]["status"] == "UNRESOLVED":
            passport["intake_questions"].append(
                {"field": field, "question_vi": f"Vui lòng xác nhận {label}.", "question_en": f"Please confirm {field}."}
            )
    passport["audit_log"].append(
        {"action": "INTAKE_EXISTING_PROJECT", "artifact": str(project.path), "status": "DRAFT_INFERRED"}
    )
    return passport
