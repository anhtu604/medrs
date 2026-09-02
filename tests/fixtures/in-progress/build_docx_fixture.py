from pathlib import Path

from docx import Document


def build(path: Path) -> Path:
    document = Document()
    document.add_heading("Đánh giá phơi nhiễm X và kết cục Y", level=0)
    document.add_heading("Đặt vấn đề", level=1)
    document.add_paragraph("Mục tiêu nghiên cứu: mô tả tỷ lệ kết cục Y ở người bệnh có phơi nhiễm X.")
    document.add_heading("Đối tượng và phương pháp", level=1)
    document.add_paragraph("Đối tượng nghiên cứu là người bệnh trưởng thành điều trị tại bệnh viện A.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Biến"
    table.cell(0, 1).text = "Định nghĩa"
    table.cell(1, 0).text = "Kết cục Y"
    table.cell(1, 1).text = "Có/không"
    document.add_heading("Mục tiêu nghiên cứu", level=1)
    document.add_paragraph("Phân tích mối liên quan giữa phơi nhiễm X và kết cục Y.")
    document.add_paragraph("Tài liệu tham khảo: Nguyen A. Example study. 2024.")
    document.save(path)
    return path


if __name__ == "__main__":
    build(Path(__file__).with_name("thesis-partial.docx"))

