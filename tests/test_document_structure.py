from pathlib import Path

from docx import Document

from medical_research_skills_vn.structure_profiles import (
    apply_docx_structure,
    inspect_docx_structure,
    load_structure_profile,
    plan_restructure,
)


ROOT = Path(__file__).parents[1]
PROFILE = ROOT / "profiles/institution/hmu/thesis-master-2020-current-2026.yaml"


def _make_misordered_docx(path: Path):
    document = Document()
    document.add_heading("KẾT LUẬN", level=1)
    document.add_paragraph("Nội dung kết luận.")
    document.add_heading("CHƯƠNG 3. KẾT QUẢ", level=1)
    document.add_paragraph("Nội dung kết quả.")
    document.add_heading("ĐẶT VẤN ĐỀ", level=1)
    document.add_paragraph("Nội dung đặt vấn đề.")
    document.add_heading("CHƯƠNG 2. ĐỐI TƯỢNG VÀ PHƯƠNG PHÁP NGHIÊN CỨU", level=1)
    document.add_paragraph("Nội dung phương pháp.")
    document.add_heading("CHƯƠNG 4. BÀN LUẬN", level=1)
    document.add_paragraph("Nội dung bàn luận.")
    document.add_heading("CHƯƠNG 1. TỔNG QUAN", level=1)
    document.add_paragraph("Nội dung tổng quan.")
    document.add_heading("TÀI LIỆU THAM KHẢO", level=1)
    document.add_paragraph("Một tài liệu.")
    document.save(path)


def test_restructure_requires_author_approval(tmp_path):
    source = tmp_path / "draft.docx"
    _make_misordered_docx(source)
    profile = load_structure_profile(PROFILE)

    plan = plan_restructure(inspect_docx_structure(source, profile)["recognized"], profile)
    assert plan["status"] == "AUTHOR_APPROVAL_REQUIRED"
    assert plan["moves"]


def test_reorders_actual_docx_and_validates_generated_artifact(tmp_path):
    source = tmp_path / "draft.docx"
    output = tmp_path / "draft-structured.docx"
    _make_misordered_docx(source)
    profile = load_structure_profile(PROFILE)

    result = apply_docx_structure(source, output, profile, author_approved=True)

    assert result["status"] == "STRUCTURE_APPLIED"
    assert output.exists()
    assert result["validation"]["status"] == "PASS"
    assert result["validation"]["recognized"] == [
        "introduction",
        "chapter-1-literature-review",
        "chapter-2-methods",
        "chapter-3-results",
        "chapter-4-discussion",
        "conclusion",
        "references",
    ]
    assert source.read_bytes() != output.read_bytes()


def test_never_overwrites_source_or_invents_missing_sections(tmp_path):
    source = tmp_path / "draft.docx"
    _make_misordered_docx(source)
    profile = load_structure_profile(PROFILE)

    refused = apply_docx_structure(source, source, profile, author_approved=True)
    assert refused["status"] == "SOURCE_OVERWRITE_FORBIDDEN"

    doc = Document(source)
    for paragraph in list(doc.paragraphs):
        if "BÀN LUẬN" in paragraph.text:
            paragraph._element.getparent().remove(paragraph._element)
    missing_source = tmp_path / "missing.docx"
    doc.save(missing_source)
    missing_output = tmp_path / "missing-structured.docx"
    result = apply_docx_structure(missing_source, missing_output, profile, author_approved=True)
    assert result["status"] == "AUTHOR_INPUT_REQUIRED"
    assert "chapter-4-discussion" in result["missing_required"]
    assert not missing_output.exists()


def test_duplicate_semantic_section_is_not_silently_dropped(tmp_path):
    source = tmp_path / "duplicate.docx"
    _make_misordered_docx(source)
    document = Document(source)
    document.add_heading("CHƯƠNG 3. KẾT QUẢ", level=1)
    document.add_paragraph("Một bản kết quả thứ hai.")
    document.save(source)
    profile = load_structure_profile(PROFILE)

    output = tmp_path / "duplicate-structured.docx"
    result = apply_docx_structure(source, output, profile, author_approved=True)
    assert result["status"] == "AUTHOR_INPUT_REQUIRED"
    assert result["duplicate_sections"] == ["chapter-3-results"]
    assert not output.exists()
