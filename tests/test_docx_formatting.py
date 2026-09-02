from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement

from medical_research_skills_vn.docx_formatting import format_docx
from medical_research_skills_vn.docx_validation import validate_docx
from medical_research_skills_vn.structure_profiles import load_structure_profile


ROOT = Path(__file__).parents[1]
PROFILE_PATH = ROOT / "profiles/institution/hmu/word-format-master-2020-current-2026.yaml"


def _fixture(path: Path):
    document = Document()
    document.styles.add_style("CustomBody", WD_STYLE_TYPE.PARAGRAPH)
    document.add_heading("MỤC LỤC", level=1)
    document.add_paragraph("Mục lục cũ")
    document.add_heading("ĐẶT VẤN ĐỀ", level=1)
    body = document.add_paragraph("Nội dung thử nghiệm có chỉnh sửa.", style="CustomBody")
    insertion = OxmlElement("w:ins")
    insertion.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id", "7")
    body._p.append(insertion)
    document.add_heading("CHƯƠNG 1. TỔNG QUAN", level=1)
    document.add_paragraph("Tổng quan.")
    document.save(path)


def test_ooxml_formatter_never_overwrites_source(tmp_path):
    source = tmp_path / "source.docx"
    _fixture(source)
    profile = load_structure_profile(PROFILE_PATH)
    result = format_docx(source, source, profile, author_approved=True)
    assert result["status"] == "SOURCE_OVERWRITE_FORBIDDEN"


def test_formats_and_validates_the_generated_docx_not_configuration(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "formatted.docx"
    _fixture(source)
    profile = load_structure_profile(PROFILE_PATH)

    result = format_docx(source, output, profile, author_approved=True)
    report = validate_docx(output, profile, rendered_artifacts=[])

    assert result["status"] == "FORMATTED_WITH_UNPERFORMED_CHECKS"
    assert output.exists()
    assert report["artifact_hash"] == result["output_hash"]
    assert report["mechanical_checks"]["normal_style"] == "PASS"
    assert report["mechanical_checks"]["direct_formatting"] == "PASS"
    assert report["mechanical_checks"]["margins"] == "PASS"
    assert report["mechanical_checks"]["toc_field"] == "PASS"
    assert report["mechanical_checks"]["page_number_field"] == "PASS"
    assert report["mechanical_checks"]["page_restart_at_introduction"] == "PASS"
    assert report["rendered_checks"]["pagination"] == "AUTHOR_APPROVAL_REQUIRED"
    assert report["rendered_checks"]["visual_layout"] == "AUTHOR_APPROVAL_REQUIRED"


def test_preserves_tracked_change_xml_and_creates_new_file(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "formatted.docx"
    _fixture(source)
    profile = load_structure_profile(PROFILE_PATH)
    before = source.read_bytes()

    format_docx(source, output, profile, author_approved=True)
    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml")
    assert b"<w:ins" in xml
    assert source.read_bytes() == before


def test_formatting_requires_profile_and_author_approval(tmp_path):
    source = tmp_path / "source.docx"
    output = tmp_path / "formatted.docx"
    _fixture(source)
    profile = load_structure_profile(PROFILE_PATH)
    assert format_docx(source, output, profile, author_approved=False)["status"] == "AUTHOR_APPROVAL_REQUIRED"
    stale = dict(profile, verification_status="STALE")
    assert format_docx(source, output, stale, author_approved=True)["status"] == "OFFICIAL_RULE_REQUIRED"
